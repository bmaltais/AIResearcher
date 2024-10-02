import os
from dotenv import load_dotenv
import hashlib
import argparse
from typing import List
import PyPDF2
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import email
import markdown
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.llms import Ollama
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from tqdm import tqdm
import html2text
import re
from langchain.chains.summarize import load_summarize_chain

load_dotenv()

# Configuration
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 0
MAX_CHUNK_SIZE = 1024  
MODEL = "nomic-embed-text"
SUMMARY_MODEL = "llama3.1"  # You can change this to your preferred local model
LANGCHAIN_TRACING_V2 = os.getenv('LANGCHAIN_TRACING_V2')
LANGCHAIN_API_KEY = os.getenv('LANGCHAIN_API_KEY')

# Initialize Ollama embeddings
embeddings = OllamaEmbeddings(model=MODEL)

def parse_arguments():
    parser = argparse.ArgumentParser(
        description="Process documents and store in ChromaDB"
    )
    parser.add_argument(
        "--collection",
        type=str,
        default="default_collection",
        help="Name of the ChromaDB collection",
    )
    parser.add_argument(
        "--document-directory",
        type=str,
        default="./documents",
        help="Directory containing documents to process",
    )
    parser.add_argument(
        "--chroma-db-path",
        type=str,
        default="./chroma_db",
        help="Path to store ChromaDB",
    )
    parser.add_argument(
        "--summarize",
        action="store_true",
        help="Enable document summarization",
    )
    return parser.parse_args()

def calculate_file_hash(file_path: str) -> str:
    with open(file_path, "rb") as file:
        return hashlib.md5(file.read()).hexdigest()

def document_exists(db, file_path: str) -> bool:
    file_hash = calculate_file_hash(file_path)
    results = db.get(where={"source": file_path})
    return any(
        metadata.get("file_hash") == file_hash
        for metadata in results.get("metadatas", [])
    )

def extract_text_from_pdf(file_path):
    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        return "\n".join(page.extract_text() for page in reader.pages)

def extract_text_from_epub(file_path):
    book = epub.read_epub(file_path)
    text = ""
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            soup = BeautifulSoup(item.get_content(), "html.parser")
            text += soup.get_text()
    return text

def extract_text_from_markdown(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        md_text = file.read()
    html = markdown.markdown(md_text)
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text()

def extract_text_from_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()

def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def extract_text_from_mht(file_path):
    with open(file_path, "rb") as file:
        msg = email.message_from_binary_file(file)
    for part in msg.walk():
        if part.get_content_type() == "text/html":
            html_content = part.get_payload(decode=True).decode()
            return html2text.html2text(html_content)
    return ""

def extract_text(file_path: str) -> str:
    _, ext = os.path.splitext(file_path)
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".epub":
        return extract_text_from_epub(file_path)
    elif ext == ".md":
        return extract_text_from_markdown(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".mht":
        return extract_text_from_mht(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def should_skip_file(file_name):
    return file_name.endswith('-summary.txt')

def chunk_text(text: str) -> List[Document]:
    def split_into_sentences(text):
        return re.split(r'(?<=[.!?])\s+', text)

    sentences = split_into_sentences(text)
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= CHUNK_SIZE:
            current_chunk += sentence + " "
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            if len(sentence) > MAX_CHUNK_SIZE:
                # If a single sentence is too long, split it into smaller chunks
                sentence_chunks = [sentence[i:i+MAX_CHUNK_SIZE] for i in range(0, len(sentence), MAX_CHUNK_SIZE)]
                chunks.extend(sentence_chunks)
            else:
                current_chunk = sentence + " "

    if current_chunk:
        chunks.append(current_chunk.strip())

    # Create overlapping chunks
    overlapped_chunks = []
    for i in range(len(chunks)):
        start = max(0, i - 1)
        end = min(len(chunks), i + 2)
        overlapped_chunk = " ".join(chunks[start:end])
        overlapped_chunks.append(overlapped_chunk)

    return [Document(page_content=chunk) for chunk in overlapped_chunks]

def setup_local_llm(model_name: str):
    llm = Ollama(model=model_name, temperature=0)
    return llm

# def summarize_chunk(llm, chunk: str) -> str:
#     prompt = PromptTemplate(
#         input_variables=["chunk"],
#         template="Create a cohesive and shortened version of about 300 words from the following text:\n\n{chunk} Write it as if you were that author in the same style of the original text, not as an expressed summary."
#     )
#     chain = LLMChain(llm=llm, prompt=prompt)
#     chunk_summary = chain.run(chunk=chunk)
#     print(chunk_summary)
#     return chain.run(chunk=chunk)

def summarize_document(llm, text) -> str:
    
    print("Creating final summary...")
    
    split_docs = chunk_text(text)
    prompt_template = """
    Write a news article based on the following information. Ensure the article is clear, informative, and captures the main points of the provided text. Maintain a professional and neutral tone as typical of news reporting:

    {text}

    NEWS ARTICLE:
    """
    prompt = PromptTemplate.from_template(prompt_template)

    refine_template = """
    Your job is to enhance an existing news article by adding more details based on the new context provided.
    Ensure the article remains clear, informative, and maintains a professional and neutral tone.
    Do not explain what you added; simply update the article and return it.

    We have provided an existing news article up to a certain point:
    ************
    {existing_answer}
    ************
    We have the opportunity to refine the existing article (only if needed) with some more context below.
    ------------
    {text}
    ------------
    Given the new context, refine the original news article in English. If the context isn't useful, return the original article.
    """

    # prompt_template = """
    # Write a concise summary of the following as if you were the author of the document. Keep the same style at the original document:

    # {text}

    # CONCISE SUMMARY:
    # """
    # prompt = PromptTemplate.from_template(prompt_template)

    # refine_template = """
    # Your job is to add details to an existing summary as if you were the author of the document.
    # Keep the same writing style as the document. Do not explain what you added. Just update the summary and return it without emiting an opinion about it.
    # We have provided an existing summary up to a certain point:
    # ************
    # {existing_answer}
    # ************
    # We have the opportunity to refine the existing summary (only if needed) with some more context below.
    # ------------
    # {text}
    # ------------
    # Given the new context, refine the original summary in English. If the context isn't useful, return the original summary.
    # """
    refine_prompt = PromptTemplate.from_template(refine_template)
    chain = load_summarize_chain(
        llm=llm,
        chain_type="refine",
        question_prompt=prompt,
        refine_prompt=refine_prompt,
        return_intermediate_steps=True,
        input_key="input_documents",
        output_key="output_text",
    )
    result = chain.invoke({"input_documents": split_docs}, return_only_outputs=True)

    print("Text summary:\n\n")
    print(result["output_text"])
    return result["output_text"]

def process_file(db, file_path: str, summarize: bool = False, llm=None):
    if document_exists(db, file_path):
        print(f"Document {file_path} already exists in the database. Skipping.")
        return

    print(f"Processing {os.path.basename(file_path)}")
    text = extract_text(file_path)

    # Create ./tmp/ folder if it doesn't exist
    tmp_folder = "./tmp/"
    os.makedirs(tmp_folder, exist_ok=True)

    # Get the base name of the file and replace the extension with .txt
    base_name = os.path.basename(file_path)
    new_file_name = os.path.splitext(base_name)[0] + ".txt"
    new_file_path = os.path.join(tmp_folder, new_file_name)

    # Write the content to the new file
    with open(new_file_path, "w", encoding="utf-8") as f:
        f.write(text)

    print(f"Wrote content to {new_file_path}")

    chunks = chunk_text(text)
    file_hash = calculate_file_hash(file_path)

    if summarize and llm:
        print(f"Summarizing {os.path.basename(file_path)}")
        summary = summarize_document(llm, text)
        
        # Save the summary to a new file
        summary_file_name = os.path.splitext(base_name)[0] + "-summary.txt"
        summary_file_path = os.path.join(os.path.dirname(file_path), summary_file_name)
        
        with open(summary_file_path, "w", encoding="utf-8") as f:
            f.write(summary)
        
        print(f"Saved summary to {summary_file_path}")

        summary_chunk = Document(
            page_content=summary,
            metadata={
                "file_hash": file_hash,
                "source": file_path,
                "chunk": "summary",
            }
        )
        chunks.append(summary_chunk)

    updated_chunks = [
        Document(
            page_content=chunk.page_content,
            metadata={
                **chunk.metadata,
                "file_hash": file_hash,
                "source": file_path,
                "chunk": i,
            },
        )
        for i, chunk in enumerate(chunks)
    ]

    print(f"Adding chunks to database...")
    db.add_documents(updated_chunks)
    print(f"Processed {os.path.basename(file_path)}")

def cleanup_chromadb(db, document_directory: str):
    all_docs = db.get()
    existing_files = {
        os.path.abspath(os.path.join(root, file))
        for root, _, files in os.walk(document_directory)
        for file in files
    }

    docs_to_remove = [
        doc_id
        for doc_id, metadata in zip(all_docs["ids"], all_docs["metadatas"])
        if "source" in metadata
        and os.path.abspath(metadata["source"]) not in existing_files
    ]

    if docs_to_remove:
        db.delete(ids=docs_to_remove)
        print(
            f"Removed {len(docs_to_remove)} documents from ChromaDB that no longer exist in the file system."
        )
    else:
        print("No documents to remove from ChromaDB.")

    print(f"Total documents remaining in ChromaDB: {len(db.get()['ids'])}")

def process_documents_to_chroma(db, directory_path: str, summarize: bool = False):
    llm = setup_local_llm(SUMMARY_MODEL) if summarize else None
    
    files_to_process = []
    for root, _, files in os.walk(directory_path):
        for file in files:
            if not should_skip_file(file):
                files_to_process.append(os.path.join(root, file))
    
    for file_path in tqdm(files_to_process, desc="Processing files", unit="file"):
        process_file(db, file_path, summarize, llm)

    cleanup_chromadb(db, directory_path)
    print("All documents have been processed and inserted into ChromaDB.")

if __name__ == "__main__":
    args = parse_arguments()

    # Initialize ChromaDB with the specified collection
    db = Chroma(
        persist_directory=args.chroma_db_path,
        embedding_function=embeddings,
        collection_name=args.collection,
    )

    process_documents_to_chroma(db, args.document_directory, args.summarize)